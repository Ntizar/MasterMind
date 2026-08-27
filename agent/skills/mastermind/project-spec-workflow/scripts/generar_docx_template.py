#!/usr/bin/env python3
"""
Generar DOCX profesional con tablas estilizadas.

Uso: python3 generar_docx_template.py <output.docx>

Genera un DOCX con:
- Encabezados con colores (#2563eb azul, #f97316 naranja)
- Tablas con encabezados azules y filas alternas
- Estilos de párrafo consistentes
- Footer con atribución

Se puede usar como plantilla para generar DOCX de cualquier proyecto.
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BLUE = RGBColor(0x25, 0x63, 0xEB)
ORANGE = RGBColor(0xF9, 0x73, 0x16)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
LIGHT_BLUE_BG = 'F0F4FF'
WHITE_BG = 'FFFFFF'


def apply_style(doc):
    """Aplicar estilos globales al documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def add_styled_table(doc, headers, rows):
    """Crear tabla con estilo profesional (azul encabezado, filas alternas, bordes)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2563EB',
            qn('w:val'): 'clear',
        })
        shading.append(shd)

    # Filas de datos
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            bg = WHITE_BG if row_idx % 2 == 0 else LIGHT_BLUE_BG
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): bg,
                qn('w:val'): 'clear',
            })
            shading.append(shd)

    # Bordes
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = tbl.makeelement(qn('w:tblBorders'), {})
        tblPr.insert(0, borders)
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): 'CCCCCC',
        })
        borders.append(border)

    return table


if __name__ == '__main__':
    output_path = sys.argv[1] if len(sys.argv) > 1 else '/tmp/template.docx'

    doc = Document()
    apply_style(doc)

    # Ejemplo: tabla de datos
    add_styled_table(doc,
        ['Campo', 'Valor'],
        [
            ['Nombre', 'Ejemplo de dato 1'],
            ['Descripción', 'Ejemplo de dato 2'],
        ]
    )

    doc.save(output_path)
    print(f'✅ DOCX generado: {output_path}')
